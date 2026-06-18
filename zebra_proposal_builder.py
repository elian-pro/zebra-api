"""
zebra_proposal_builder.py

Constructor de propuestas estratégicas de Zebra High Performance Marketing.
Genera documentos DOCX con el Look and Feel oficial de Zebra.

Uso:
    from zebra_proposal_builder import generate_proposal
    generate_proposal(proposal_data, output_path="propuesta.docx")

El esquema de `proposal_data` está documentado en la docstring de generate_proposal().
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# =============================================================================
# PALETA OFICIAL ZEBRA
# =============================================================================
ZEBRA_YELLOW = "F9D626"       # Acento principal
ZEBRA_BLACK = "111111"        # Texto principal / fondos oscuros
ZEBRA_DARK_GREY = "333333"    # Texto secundario
ZEBRA_MID_GREY = "888888"     # Texto terciario / detalles
ZEBRA_LIGHT_GREY = "DDDDDD"   # Bordes / separadores suaves
ZEBRA_BG_LIGHT = "F5F5F5"     # Fondo de cajas neutras
ZEBRA_WHITE = "FFFFFF"        # Texto sobre fondos oscuros

# Fuente oficial
ZEBRA_FONT = "Arial"


# =============================================================================
# HELPERS DE FORMATO
# =============================================================================

def set_cell_background(cell, color_hex):
    """Establece el color de fondo de una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4,
                     top=True, bottom=True, left=True, right=True):
    """Aplica bordes a una celda. Size en octavos de punto (4 = 0.5pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')

    for edge, enabled in [('top', top), ('bottom', bottom),
                           ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{edge}')
        if enabled:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), color)
        else:
            border.set(qn('w:val'), 'nil')
        tc_borders.append(border)

    tc_pr.append(tc_borders)


def remove_cell_borders(cell):
    """Quita todos los bordes visibles de una celda."""
    set_cell_borders(cell, top=False, bottom=False, left=False, right=False)


def set_row_cant_split(row):
    """
    Evita que una fila de tabla se rompa entre páginas.
    Si la fila no cabe completa en la página actual, se mueve completa a la siguiente.
    """
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_table_cant_split(table):
    """Aplica cantSplit a todas las filas de una tabla."""
    for row in table.rows:
        set_row_cant_split(row)


def set_keep_with_next(paragraph):
    """
    Mantiene un párrafo junto al siguiente. Útil para que un encabezado
    no se separe del contenido que le sigue.
    """
    pf = paragraph.paragraph_format
    pf.keep_with_next = True


def set_cell_vertical_alignment(cell, alignment='center'):
    """center, top, bottom"""
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), alignment)
    tc_pr.append(v_align)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Margenes internos en twentieths of a point (1/20 pt). 120 ≈ 0.08 in."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def style_run(run, size=10, bold=False, color=ZEBRA_DARK_GREY,
              font=ZEBRA_FONT, italic=False):
    """Aplica estilo a un run de texto."""
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    # Asegurar que se aplique la fuente Arial en el XML
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn('w:ascii'), font)
    r_fonts.set(qn('w:hAnsi'), font)
    r_fonts.set(qn('w:cs'), font)


def add_styled_paragraph(container, text, size=10, bold=False,
                         color=ZEBRA_DARK_GREY, alignment=None,
                         space_before=0, space_after=0, line_spacing=1.15,
                         italic=False):
    """
    Agrega un párrafo estilizado a un contenedor (doc o cell).
    Retorna el párrafo creado.
    """
    p = container.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_multi_run_paragraph(container, runs, alignment=None,
                             space_before=0, space_after=0, line_spacing=1.15):
    """
    Agrega un párrafo con múltiples runs de formato distinto.
    `runs` es una lista de dicts con keys: text, size, bold, color.
    """
    p = container.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    for run_spec in runs:
        r = p.add_run(run_spec.get('text', ''))
        style_run(
            r,
            size=run_spec.get('size', 10),
            bold=run_spec.get('bold', False),
            color=run_spec.get('color', ZEBRA_DARK_GREY),
            italic=run_spec.get('italic', False)
        )
    return p


def set_page_setup(doc):
    """Carta, márgenes 0.75 pulgadas."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def set_default_document_font(doc, font=ZEBRA_FONT):
    """Establece Arial como fuente default del documento."""
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = font
    normal.font.size = Pt(10)
    # Asegurar que se aplique a ASCII y CS
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn('w:ascii'), font)
    r_fonts.set(qn('w:hAnsi'), font)
    r_fonts.set(qn('w:cs'), font)


def add_spacer(doc, height_pt=6):
    """Agrega un párrafo vacío como espaciador vertical."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)


# =============================================================================
# BLOQUES CONSTRUCTIVOS DE LA PROPUESTA
# =============================================================================

def add_wordmark_logo(paragraph, text="ZEBRA", size=20, color=ZEBRA_BLACK,
                      tracking=120):
    """
    Renderiza el wordmark "ZEBRA" como logo tipográfico con tracking
    (120 = 6 pt en veinteavos de punto) para reproducir el look oficial.
    """
    r = paragraph.add_run(text)
    style_run(r, size=size, bold=True, color=color)
    r_pr = r._element.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(tracking))
    r_pr.append(spacing)
    return r


def build_title_block(doc, titulo="PROPUESTA ESTRATÉGICA", subtitulo=""):
    """
    Título principal con logo ZEBRA en la esquina superior derecha.
    Izquierda: títulos. Derecha: wordmark "ZEBRA".
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_width = Inches(5.0)
    right_width = Inches(2.0)
    table.columns[0].width = left_width
    table.columns[1].width = right_width

    row = table.rows[0]
    left, right = row.cells[0], row.cells[1]
    left.width = left_width
    right.width = right_width

    # Columna izquierda: títulos
    set_cell_vertical_alignment(left, 'center')
    left.paragraphs[0].clear()
    p_titulo = left.paragraphs[0]
    pf = p_titulo.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    r = p_titulo.add_run(titulo)
    style_run(r, size=10, bold=True, color=ZEBRA_MID_GREY)
    if subtitulo:
        add_styled_paragraph(
            left, subtitulo,
            size=22, bold=True, color=ZEBRA_BLACK,
            space_before=0, space_after=0, line_spacing=1.0
        )

    # Columna derecha: logo ZEBRA (esquina superior derecha)
    set_cell_vertical_alignment(right, 'top')
    right.paragraphs[0].clear()
    p_logo = right.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    plf = p_logo.paragraph_format
    plf.space_before = Pt(0)
    plf.space_after = Pt(0)
    add_wordmark_logo(p_logo, "ZEBRA", size=20, color=ZEBRA_BLACK)

    add_spacer(doc, 8)


def build_header_tripartita(doc, cliente, modelo, objetivo):
    """
    Tabla de 3 columnas: CLIENTE | MODELO | OBJETIVO
    Estilo: fondo claro, texto oscuro, etiquetas pequeñas arriba.
    """
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_cant_split(table)

    # Anchos totales = 9360 dxa (8.5" - 1.5" márgenes = 7")
    col_width = Inches(7.0 / 3)
    for col in table.columns:
        col.width = col_width

    row = table.rows[0]

    labels_values = [
        ("CLIENTE", cliente),
        ("MODELO", modelo),
        ("OBJETIVO", objetivo),
    ]

    for i, (label, value) in enumerate(labels_values):
        cell = row.cells[i]
        cell.width = col_width
        set_cell_background(cell, ZEBRA_BG_LIGHT)
        set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        set_cell_vertical_alignment(cell, 'center')

        # Limpiar primer párrafo
        cell.paragraphs[0].clear()
        p_label = cell.paragraphs[0]
        pf = p_label.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        r = p_label.add_run(label)
        style_run(r, size=8, bold=True, color=ZEBRA_MID_GREY)

        add_styled_paragraph(
            cell, value,
            size=10, bold=True, color=ZEBRA_BLACK,
            space_before=0, space_after=0, line_spacing=1.2
        )

    add_spacer(doc, 10)


def build_section_header(doc, numero, titulo):
    """
    Encabezado de sección: barra horizontal completa.
    Izquierda: cuadrito amarillo con el número en negro.
    Derecha: barra negra con el título en blanco.
    Ej: [01] CONTEXTO DEL CLIENTE
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(0.6)
    table.columns[1].width = Inches(6.4)
    set_table_cant_split(table)

    # Celda del número (amarilla)
    cell_num = table.rows[0].cells[0]
    cell_num.width = Inches(0.6)
    set_cell_background(cell_num, ZEBRA_YELLOW)
    set_cell_borders(cell_num, color=ZEBRA_YELLOW, size=4)
    set_cell_margins(cell_num, top=120, bottom=120, left=0, right=0)
    set_cell_vertical_alignment(cell_num, 'center')

    cell_num.paragraphs[0].clear()
    p = cell_num.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run(str(numero))
    style_run(r, size=11, bold=True, color=ZEBRA_BLACK)

    # Celda del título (negra)
    cell_title = table.rows[0].cells[1]
    cell_title.width = Inches(6.4)
    set_cell_background(cell_title, ZEBRA_BLACK)
    set_cell_borders(cell_title, color=ZEBRA_BLACK, size=4)
    set_cell_margins(cell_title, top=120, bottom=120, left=220, right=220)
    set_cell_vertical_alignment(cell_title, 'center')

    cell_title.paragraphs[0].clear()
    p = cell_title.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run(titulo)
    style_run(r, size=11, bold=True, color=ZEBRA_WHITE)

    add_spacer(doc, 6)


def build_context_section(doc, intro, proyectos_titulo, proyectos_items,
                           situacion_titulo, situacion_texto):
    """
    Sección 01 — Contexto del cliente.
    Párrafo intro + tabla de 2 columnas (proyectos / situación actual).
    """
    add_styled_paragraph(
        doc, intro,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=10, line_spacing=1.4
    )

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_width = Inches(3.5)
    for col in table.columns:
        col.width = col_width
    set_table_cant_split(table)

    # Celda 1 — Proyectos
    cell_left = table.rows[0].cells[0]
    cell_left.width = col_width
    set_cell_background(cell_left, ZEBRA_BG_LIGHT)
    set_cell_borders(cell_left, color=ZEBRA_LIGHT_GREY, size=4)
    set_cell_margins(cell_left, top=160, bottom=160, left=180, right=180)

    cell_left.paragraphs[0].clear()
    p = cell_left.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    r = p.add_run(proyectos_titulo)
    style_run(r, size=8, bold=True, color=ZEBRA_MID_GREY)

    for item in proyectos_items:
        add_styled_paragraph(
            cell_left, item,
            size=10, bold=True, color=ZEBRA_BLACK,
            space_before=0, space_after=2
        )

    # Celda 2 — Situación actual
    cell_right = table.rows[0].cells[1]
    cell_right.width = col_width
    set_cell_background(cell_right, ZEBRA_BG_LIGHT)
    set_cell_borders(cell_right, color=ZEBRA_LIGHT_GREY, size=4)
    set_cell_margins(cell_right, top=160, bottom=160, left=180, right=180)

    cell_right.paragraphs[0].clear()
    p = cell_right.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    r = p.add_run(situacion_titulo)
    style_run(r, size=8, bold=True, color=ZEBRA_MID_GREY)

    add_styled_paragraph(
        cell_right, situacion_texto,
        size=9.5, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=0, line_spacing=1.4
    )

    add_spacer(doc, 8)


def build_diagnosis_boxes(doc, diagnostics):
    """
    Sección 02 — Diagnóstico.
    Una sola caja gris por diagnóstico:
      - Título en negro bold arriba.
      - Párrafo con ⚡ amarillo al inicio + texto en gris oscuro.
    `diagnostics` es lista de dicts: {titulo, texto}
    """
    for diag in diagnostics:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(7.0)
        set_table_cant_split(table)

        cell = table.rows[0].cells[0]
        cell.width = Inches(7.0)
        set_cell_background(cell, ZEBRA_BG_LIGHT)
        set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cell, top=160, bottom=160, left=220, right=220)

        # Título en bold negro
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        r = p.add_run(diag['titulo'])
        style_run(r, size=10.5, bold=True, color=ZEBRA_BLACK)

        # Texto con ⚡ amarillo al inicio
        add_multi_run_paragraph(
            cell,
            runs=[
                {'text': '⚡  ', 'size': 11, 'bold': True, 'color': ZEBRA_YELLOW},
                {'text': diag['texto'], 'size': 9.5, 'color': ZEBRA_DARK_GREY},
            ],
            space_before=0, space_after=0, line_spacing=1.4
        )

        add_spacer(doc, 4)

    add_spacer(doc, 4)


def build_paradigm_shift(doc, antes_titulo, antes_texto,
                          despues_titulo, despues_texto, punchline):
    """
    Sección 03 — Cambio de paradigma.
    Dos bloques secuenciales + caja de punchline en amarillo.
    """
    # ANTES
    add_styled_paragraph(
        doc, antes_titulo,
        size=9, bold=True, color=ZEBRA_MID_GREY,
        space_before=0, space_after=2
    )
    add_styled_paragraph(
        doc, antes_texto,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=10, line_spacing=1.4
    )

    # DESPUÉS
    add_styled_paragraph(
        doc, despues_titulo,
        size=9, bold=True, color=ZEBRA_MID_GREY,
        space_before=0, space_after=2
    )
    add_styled_paragraph(
        doc, despues_texto,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=12, line_spacing=1.4
    )

    # Caja punchline - NEGRA con texto amarillo
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    set_table_cant_split(table)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, ZEBRA_BLACK)
    set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
    set_cell_margins(cell, top=180, bottom=180, left=220, right=220)

    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.4
    r = p.add_run(punchline)
    style_run(r, size=11, bold=True, color=ZEBRA_YELLOW)

    add_spacer(doc, 8)


def build_system_models(doc, intro, modelos):
    """
    Sección 04 — Sistema propuesto.
    Intro + fila de N cajas.
    Primer modelo (dominante): caja NEGRA con título amarillo + texto blanco.
    Modelos de soporte: cajas GRISES con título negro bold + texto negro regular.
    `modelos` es lista de {nombre, descripcion}
    """
    add_styled_paragraph(
        doc, intro,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=10, line_spacing=1.4
    )

    n = len(modelos)
    if n == 0:
        return

    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    col_width = Inches(7.0 / n)
    for col in table.columns:
        col.width = col_width
    set_table_cant_split(table)

    for i, modelo in enumerate(modelos):
        cell = table.rows[0].cells[i]
        cell.width = col_width

        if i == 0:
            # Modelo dominante — caja negra
            set_cell_background(cell, ZEBRA_BLACK)
            set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
            title_color = ZEBRA_YELLOW
            text_color = ZEBRA_WHITE
        else:
            # Modelos de soporte — cajas grises
            set_cell_background(cell, ZEBRA_BG_LIGHT)
            set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4)
            title_color = ZEBRA_BLACK
            text_color = ZEBRA_DARK_GREY

        set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
        set_cell_vertical_alignment(cell, 'top')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        r = p.add_run(modelo['nombre'])
        style_run(r, size=11, bold=True, color=title_color)

        add_styled_paragraph(
            cell, modelo['descripcion'],
            size=9, color=text_color,
            space_before=0, space_after=0, line_spacing=1.4
        )

    add_spacer(doc, 8)


def build_architecture(doc, capas):
    """
    Sección 05 — Arquitectura del sistema.
    Cada capa tiene nombre (como texto simple encima) + items numerados con bullets.
    `capas` es lista de {nombre, items: [{numero, titulo, bullets}]}
    """
    for capa in capas:
        # Encabezado de capa: texto simple en gris, sin caja oscura
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(6)
        r = p.add_run(capa['nombre'])
        style_run(r, size=9, bold=True, color=ZEBRA_MID_GREY)

        for item in capa['items']:
            _build_arch_item(doc, item)

        add_spacer(doc, 4)


def build_architecture_flat(doc, items):
    """
    Variante sin capas: solo items numerados.
    """
    for item in items:
        _build_arch_item(doc, item)
    add_spacer(doc, 4)


def _build_arch_item(doc, item):
    """Un item numerado de arquitectura:
    - Número en caja NEGRA a la izquierda (número grande amarillo).
    - Contenido en caja GRIS a la derecha (título negro bold + bullets).
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(6.45)
    set_table_cant_split(table)

    # Celda del número — NEGRA con número amarillo
    cell_num = table.rows[0].cells[0]
    cell_num.width = Inches(0.55)
    set_cell_background(cell_num, ZEBRA_BLACK)
    set_cell_borders(cell_num, color=ZEBRA_BLACK, size=4)
    set_cell_margins(cell_num, top=120, bottom=120, left=0, right=0)
    set_cell_vertical_alignment(cell_num, 'center')

    cell_num.paragraphs[0].clear()
    p = cell_num.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run(str(item['numero']))
    style_run(r, size=20, bold=True, color=ZEBRA_YELLOW)

    # Celda del contenido — GRIS
    cell_body = table.rows[0].cells[1]
    cell_body.width = Inches(6.45)
    set_cell_background(cell_body, ZEBRA_BG_LIGHT)
    set_cell_borders(cell_body, color=ZEBRA_LIGHT_GREY, size=4)
    set_cell_margins(cell_body, top=140, bottom=140, left=180, right=180)

    cell_body.paragraphs[0].clear()
    p = cell_body.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    r = p.add_run(item['titulo'])
    style_run(r, size=10.5, bold=True, color=ZEBRA_BLACK)

    for bullet in item['bullets']:
        add_multi_run_paragraph(
            cell_body,
            runs=[
                {'text': '→  ', 'size': 9.5, 'bold': True, 'color': ZEBRA_YELLOW},
                {'text': bullet, 'size': 9.5, 'color': ZEBRA_DARK_GREY},
            ],
            space_before=0, space_after=2, line_spacing=1.35
        )

    add_spacer(doc, 3)


def build_content_pieces(doc, intro, piezas):
    """
    Sección 06 — Piezas de contenido.
    Intro + tabla de 3 columnas: # | PIEZA | FUNCIÓN
    `piezas` es lista de {numero, pieza, funcion}
    """
    add_styled_paragraph(
        doc, intro,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=8, line_spacing=1.4
    )

    table = doc.add_table(rows=1 + len(piezas), cols=3)
    table.autofit = False
    widths = [Inches(0.5), Inches(2.5), Inches(4.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
    set_table_cant_split(table)  # ninguna fila se rompe entre páginas

    # Header
    header_texts = ["#", "PIEZA", "FUNCIÓN EN EL SISTEMA"]
    for i, text in enumerate(header_texts):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_background(cell, ZEBRA_BLACK)
        set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(text)
        style_run(r, size=8.5, bold=True, color=ZEBRA_YELLOW)

    # Filas
    for row_idx, pieza in enumerate(piezas, start=1):
        cells = table.rows[row_idx].cells

        # Número
        cells[0].width = widths[0]
        set_cell_background(cells[0], ZEBRA_WHITE)
        set_cell_borders(cells[0], color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cells[0], top=100, bottom=100, left=120, right=120)
        set_cell_vertical_alignment(cells[0], 'top')

        cells[0].paragraphs[0].clear()
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(str(pieza['numero']))
        style_run(r, size=10, bold=True, color=ZEBRA_MID_GREY)

        # Pieza
        cells[1].width = widths[1]
        set_cell_background(cells[1], ZEBRA_WHITE)
        set_cell_borders(cells[1], color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cells[1], top=100, bottom=100, left=120, right=120)
        set_cell_vertical_alignment(cells[1], 'top')

        cells[1].paragraphs[0].clear()
        p = cells[1].paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(pieza['pieza'])
        style_run(r, size=9, bold=True, color=ZEBRA_BLACK)

        # Función
        cells[2].width = widths[2]
        set_cell_background(cells[2], ZEBRA_WHITE)
        set_cell_borders(cells[2], color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cells[2], top=100, bottom=100, left=120, right=120)
        set_cell_vertical_alignment(cells[2], 'top')

        cells[2].paragraphs[0].clear()
        p = cells[2].paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(pieza['funcion'])
        style_run(r, size=9, color=ZEBRA_DARK_GREY)

    add_spacer(doc, 8)


def build_implementation_phases(doc, fases):
    """
    Sección 07 — Plan de implementación.
    4 fases, cada una con nombre, título y actividades.
    `fases` es lista de {fase, titulo, actividades}.

    NOTA: el campo `semanas` ya no se usa en el render (puede o no estar en
    el dict de entrada; si está, simplemente se ignora).
    """
    for fase in fases:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(5.5)
        set_table_cant_split(table)

        # Celda izquierda — nombre de fase (centrado vertical sin semanas debajo)
        cell_left = table.rows[0].cells[0]
        cell_left.width = Inches(1.5)
        set_cell_background(cell_left, ZEBRA_BLACK)
        set_cell_borders(cell_left, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell_left, top=140, bottom=140, left=140, right=140)
        set_cell_vertical_alignment(cell_left, 'center')

        cell_left.paragraphs[0].clear()
        p = cell_left.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(fase['fase'])
        style_run(r, size=11, bold=True, color=ZEBRA_YELLOW)

        # Celda derecha — título + actividades
        cell_right = table.rows[0].cells[1]
        cell_right.width = Inches(5.5)
        set_cell_background(cell_right, ZEBRA_BG_LIGHT)
        set_cell_borders(cell_right, color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cell_right, top=140, bottom=140, left=180, right=180)

        cell_right.paragraphs[0].clear()
        p = cell_right.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        r = p.add_run(fase['titulo'])
        style_run(r, size=10.5, bold=True, color=ZEBRA_BLACK)

        for actividad in fase['actividades']:
            add_multi_run_paragraph(
                cell_right,
                runs=[
                    {'text': '·  ', 'size': 9.5, 'bold': True, 'color': ZEBRA_YELLOW},
                    {'text': actividad, 'size': 9.5, 'color': ZEBRA_DARK_GREY},
                ],
                space_before=0, space_after=2, line_spacing=1.35
            )

        add_spacer(doc, 3)

    add_spacer(doc, 4)


def build_kpi_row(doc, kpis):
    """
    Fila de N KPIs en caja negra (2-3 típicamente).
    Números grandes en amarillo bold, descripciones en blanco.
    `kpis` es lista de {valor, descripcion}
    """
    n = len(kpis)
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    col_width = Inches(7.0 / n)
    for col in table.columns:
        col.width = col_width
    set_table_cant_split(table)

    for i, kpi in enumerate(kpis):
        cell = table.rows[0].cells[i]
        cell.width = col_width
        set_cell_background(cell, ZEBRA_BLACK)
        # Bordes negros entre columnas para que se vea como una sola caja
        set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell, top=220, bottom=220, left=120, right=120)
        set_cell_vertical_alignment(cell, 'center')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        r = p.add_run(kpi['valor'])
        style_run(r, size=22, bold=True, color=ZEBRA_YELLOW)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)
        pf2.line_spacing = 1.3
        r2 = p2.add_run(kpi['descripcion'])
        style_run(r2, size=8.5, color=ZEBRA_WHITE)

    add_spacer(doc, 8)


def build_base_scenario(doc, titulo, subtitulo, bullets, cierre):
    """
    Caja del escenario base bajo los KPIs.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    set_table_cant_split(table)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, ZEBRA_WHITE)
    set_cell_borders(
        cell,
        color=ZEBRA_YELLOW, size=8,
        top=False, bottom=False, left=True, right=False
    )
    set_cell_margins(cell, top=120, bottom=140, left=200, right=160)

    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    r = p.add_run(titulo)
    style_run(r, size=10, bold=True, color=ZEBRA_BLACK)

    add_styled_paragraph(
        cell, subtitulo,
        size=9.5, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=6, line_spacing=1.4
    )

    for bullet in bullets:
        add_multi_run_paragraph(
            cell,
            runs=[
                {'text': '·  ', 'size': 9.5, 'bold': True, 'color': ZEBRA_YELLOW},
                {'text': bullet, 'size': 9.5, 'color': ZEBRA_DARK_GREY},
            ],
            space_before=0, space_after=2, line_spacing=1.35
        )

    add_spacer(doc, 8)

    # Frase de cierre del escenario — caja NEGRA con texto amarillo
    if cierre:
        table2 = doc.add_table(rows=1, cols=1)
        table2.autofit = False
        table2.columns[0].width = Inches(7.0)
        set_table_cant_split(table2)

        cell2 = table2.rows[0].cells[0]
        cell2.width = Inches(7.0)
        set_cell_background(cell2, ZEBRA_BLACK)
        set_cell_borders(cell2, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell2, top=160, bottom=160, left=220, right=220)

        cell2.paragraphs[0].clear()
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.4
        r = p.add_run(cierre)
        style_run(r, size=10.5, bold=True, color=ZEBRA_YELLOW)

        add_spacer(doc, 8)


def build_investment_section(doc, setup_titulo, setup_monto, setup_items,
                              fee_titulo, fee_monto, fee_items,
                              pauta_texto, pauta_fase_inicial):
    """
    Sección 09 — Inversión.
    Dos cajas lado a lado (Setup / Fee) + párrafo de pauta.
    """
    # Cajas Setup + Fee
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_width = Inches(3.5)
    for col in table.columns:
        col.width = col_width
    set_table_cant_split(table)

    for i, (titulo, monto, items) in enumerate([
        (setup_titulo, setup_monto, setup_items),
        (fee_titulo, fee_monto, fee_items),
    ]):
        cell = table.rows[0].cells[i]
        cell.width = col_width
        set_cell_background(cell, ZEBRA_BLACK)
        set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
        set_cell_vertical_alignment(cell, 'top')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        r = p.add_run(titulo)
        style_run(r, size=9, bold=True, color=ZEBRA_MID_GREY)

        add_styled_paragraph(
            cell, monto,
            size=18, bold=True, color=ZEBRA_YELLOW,
            space_before=0, space_after=8
        )

        for item in items:
            add_multi_run_paragraph(
                cell,
                runs=[
                    {'text': '·  ', 'size': 9, 'bold': True, 'color': ZEBRA_YELLOW},
                    {'text': item, 'size': 9, 'color': ZEBRA_WHITE},
                ],
                space_before=0, space_after=3, line_spacing=1.35
            )

    add_spacer(doc, 8)

    # Caja de pauta
    table2 = doc.add_table(rows=1, cols=1)
    table2.autofit = False
    table2.columns[0].width = Inches(7.0)
    set_table_cant_split(table2)

    cell2 = table2.rows[0].cells[0]
    cell2.width = Inches(7.0)
    set_cell_background(cell2, ZEBRA_BG_LIGHT)
    set_cell_borders(cell2, color=ZEBRA_LIGHT_GREY, size=4)
    set_cell_margins(cell2, top=140, bottom=140, left=200, right=200)

    cell2.paragraphs[0].clear()
    p = cell2.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    r = p.add_run("PAUTA PUBLICITARIA RECOMENDADA")
    style_run(r, size=9, bold=True, color=ZEBRA_MID_GREY)

    add_styled_paragraph(
        cell2, pauta_texto,
        size=9.5, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=6, line_spacing=1.4
    )

    if pauta_fase_inicial:
        add_multi_run_paragraph(
            cell2,
            runs=[
                {'text': 'Fase inicial recomendada: ', 'size': 9.5, 'bold': True, 'color': ZEBRA_BLACK},
                {'text': pauta_fase_inicial, 'size': 9.5, 'color': ZEBRA_DARK_GREY},
            ],
            space_before=0, space_after=0, line_spacing=1.4
        )

    add_spacer(doc, 8)


def build_why_zebra(doc, agencia_normal, zebra_texto, cierre):
    """
    Sección 10 — Por qué Zebra.
    Dos cajas contrastadas (agencia normal vs Zebra) + frase de cierre.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_width = Inches(3.5)
    for col in table.columns:
        col.width = col_width
    set_table_cant_split(table)

    # Agencia normal
    cell_left = table.rows[0].cells[0]
    cell_left.width = col_width
    set_cell_background(cell_left, ZEBRA_BG_LIGHT)
    set_cell_borders(cell_left, color=ZEBRA_LIGHT_GREY, size=4)
    set_cell_margins(cell_left, top=160, bottom=160, left=180, right=180)

    cell_left.paragraphs[0].clear()
    p = cell_left.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    r = p.add_run("LO QUE HACE UNA AGENCIA NORMAL")
    style_run(r, size=8.5, bold=True, color=ZEBRA_MID_GREY)

    add_styled_paragraph(
        cell_left, agencia_normal,
        size=9.5, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=0, line_spacing=1.4
    )

    # Zebra
    cell_right = table.rows[0].cells[1]
    cell_right.width = col_width
    set_cell_background(cell_right, ZEBRA_BLACK)
    set_cell_borders(cell_right, color=ZEBRA_BLACK, size=4)
    set_cell_margins(cell_right, top=160, bottom=160, left=180, right=180)

    cell_right.paragraphs[0].clear()
    p = cell_right.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    r = p.add_run("LO QUE HACE ZEBRA")
    style_run(r, size=8.5, bold=True, color=ZEBRA_YELLOW)

    add_styled_paragraph(
        cell_right, zebra_texto,
        size=9.5, color=ZEBRA_WHITE,
        space_before=0, space_after=0, line_spacing=1.4
    )

    add_spacer(doc, 10)

    # Frase distintiva - NEGRA con texto amarillo
    if cierre:
        table2 = doc.add_table(rows=1, cols=1)
        table2.autofit = False
        table2.columns[0].width = Inches(7.0)
        set_table_cant_split(table2)

        cell2 = table2.rows[0].cells[0]
        cell2.width = Inches(7.0)
        set_cell_background(cell2, ZEBRA_BLACK)
        set_cell_borders(cell2, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell2, top=180, bottom=180, left=220, right=220)

        cell2.paragraphs[0].clear()
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.4
        r = p.add_run(cierre)
        style_run(r, size=11, bold=True, color=ZEBRA_YELLOW)

        add_spacer(doc, 8)


def build_final_closing(doc, lineas):
    """
    Sección 11 opcional — Cierre con 2-3 frases contundentes en caja destacada.
    `lineas` es lista de strings (cada una es una línea de la caja).
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    set_table_cant_split(table)

    cell = table.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, ZEBRA_BLACK)
    set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
    set_cell_margins(cell, top=220, bottom=220, left=220, right=220)

    cell.paragraphs[0].clear()
    first = True
    for linea in lineas:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.4
        r = p.add_run(linea)
        style_run(r, size=12, bold=True, color=ZEBRA_WHITE)

    add_spacer(doc, 6)


def build_investment_calculator_box(doc, calc_output, anexo_excel_nota=True):
    """
    Bloque visual condensado de la Calculadora de Inversión.
    Se inserta DENTRO de la sección 09 (Inversión), después de la pauta.

    `calc_output` es un objeto CalculadoraInversionOutput de zebra_investment_calc.
    """
    from zebra_investment_calc import fmt_mxn, fmt_pct, fmt_int, fmt_roa, fmt_short_mxn

    # Header del bloque
    add_styled_paragraph(
        doc, "CALCULADORA DE INVERSIÓN: VALOR DE PROYECTO Y ABSORCIÓN",
        size=9, bold=True, color=ZEBRA_MID_GREY,
        space_before=4, space_after=6
    )

    # Tabla de 3 filas × 3 columnas con métricas clave (caja negra estilo KPI)
    metricas_grid = [
        # Fila 1: Valor del proyecto y % inversión
        [
            ("VALOR DEL PROYECTO", fmt_short_mxn(calc_output.valor_proyecto)),
            ("% INVERSIÓN PROMEDIO", fmt_pct(calc_output.porcentaje_inversion, 1)),
            ("PRESUPUESTO TOTAL", fmt_short_mxn(calc_output.presupuesto_total)),
        ],
        # Fila 2: Operación mensual
        [
            ("PAUTA MENSUAL", fmt_short_mxn(calc_output.inversion_pauta_mensual)),
            ("LEADS MENSUALES", fmt_int(calc_output.leads_esperados_mensuales)),
            ("CPL ESTIMADO", fmt_mxn(calc_output.cpl)),
        ],
        # Fila 3: Capacidad y retorno
        [
            ("ABSORCIÓN OBJETIVO", f"{calc_output.absorcion_meses} meses"),
            ("SALA NECESARIA", f"{calc_output.sala_necesaria} asesores"),
            ("ROA ESPERADO", fmt_roa(calc_output.roa)),
        ],
    ]

    table = doc.add_table(rows=3, cols=3)
    table.autofit = False
    col_w = Inches(7.0 / 3)
    for col in table.columns:
        col.width = col_w
    set_table_cant_split(table)

    for r_idx, fila in enumerate(metricas_grid):
        for c_idx, (label, valor) in enumerate(fila):
            cell = table.rows[r_idx].cells[c_idx]
            cell.width = col_w
            set_cell_background(cell, ZEBRA_BLACK)
            set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
            set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
            set_cell_vertical_alignment(cell, 'center')

            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            r = p.add_run(label)
            style_run(r, size=7.5, bold=True, color=ZEBRA_MID_GREY)

            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf2 = p2.paragraph_format
            pf2.space_before = Pt(0)
            pf2.space_after = Pt(0)
            r2 = p2.add_run(valor)
            style_run(r2, size=14, bold=True, color=ZEBRA_YELLOW)

    add_spacer(doc, 6)

    # Cascada del funnel
    add_styled_paragraph(
        doc, "EFICIENCIA COMERCIAL ESPERADA (mensual)",
        size=9, bold=True, color=ZEBRA_MID_GREY,
        space_before=2, space_after=4
    )

    funnel_table = doc.add_table(rows=1, cols=4)
    funnel_table.autofit = False
    fcol_w = Inches(7.0 / 4)
    for col in funnel_table.columns:
        col.width = fcol_w
    set_table_cant_split(funnel_table)

    funnel_data = [
        ("CITAS", fmt_int(calc_output.citas_esperadas), fmt_pct(calc_output.tasas['cita'])),
        ("ASISTENCIAS", fmt_int(calc_output.asistencias_esperadas), fmt_pct(calc_output.tasas['asistencia'])),
        ("APARTADOS", fmt_int(calc_output.apartados_esperados), fmt_pct(calc_output.tasas['apartado'])),
        ("CIERRES", fmt_int(calc_output.cierres_esperados), fmt_pct(calc_output.tasas['cierre'])),
    ]

    for i, (label, valor, tasa) in enumerate(funnel_data):
        cell = funnel_table.rows[0].cells[i]
        cell.width = fcol_w
        set_cell_background(cell, ZEBRA_BG_LIGHT)
        set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        set_cell_vertical_alignment(cell, 'center')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        r = p.add_run(label)
        style_run(r, size=7.5, bold=True, color=ZEBRA_MID_GREY)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(2)
        r2 = p2.add_run(valor)
        style_run(r2, size=16, bold=True, color=ZEBRA_BLACK)

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf3 = p3.paragraph_format
        pf3.space_before = Pt(0)
        pf3.space_after = Pt(0)
        r3 = p3.add_run(f"({tasa})")
        style_run(r3, size=8, color=ZEBRA_MID_GREY)

    add_spacer(doc, 6)

    # Nota explicativa
    nota = (
        "Cálculo basado en ingeniería inversa: valor del proyecto × 3% (promedio fijo "
        "de inversión sobre valor) ÷ absorción objetivo. El % es punto de partida; se "
        "ajusta según ticket y dinámica del proyecto. Tasas de conversión basadas en "
        "benchmarks Zebra Real Estate."
    )
    if anexo_excel_nota:
        nota += " Detalle mensual completo en Excel anexo."

    add_styled_paragraph(
        doc, nota,
        size=8.5, italic=True, color=ZEBRA_MID_GREY,
        space_before=0, space_after=8, line_spacing=1.4
    )


def build_12_month_projection(doc, plan_output, calc_output):
    """
    Sección 10 — Proyección a 12 meses (vista trimestral compacta).

    `plan_output` es PlanInversionOutput.
    `calc_output` es CalculadoraInversionOutput (para contexto).
    """
    from zebra_investment_calc import fmt_mxn, fmt_short_mxn, fmt_int, fmt_roa

    # Intro
    intro = (
        f"Proyección de operación para 12 meses considerando crecimiento progresivo "
        f"de la sala de ventas (de {plan_output.sala_inicial} a {plan_output.sala_final} "
        f"asesores) y ventanas reales de conversión "
        f"(cita 10 días · apartado 45 días · cierre 90 días)."
    )
    add_styled_paragraph(
        doc, intro,
        size=10, color=ZEBRA_DARK_GREY,
        space_before=0, space_after=10, line_spacing=1.4
    )

    # Tabla trimestral: 5 columnas (concepto + 4 trimestres)
    headers = ["CONCEPTO", "TRIM. 1", "TRIM. 2", "TRIM. 3", "TRIM. 4"]
    filas = [
        ("Asesores promedio",
         [f"{t.asesores_promedio:.1f}" for t in plan_output.trimestres]),
        ("Leads generados",
         [fmt_int(t.leads_total) for t in plan_output.trimestres]),
        ("Inversión en pauta",
         [fmt_short_mxn(t.inversion_total) for t in plan_output.trimestres]),
        ("Citas reales",
         [fmt_int(t.citas_reales) for t in plan_output.trimestres]),
        ("Apartados reales",
         [fmt_int(t.apartados_reales) for t in plan_output.trimestres]),
        ("Cierres reales",
         [fmt_int(t.cierres_reales) for t in plan_output.trimestres]),
        ("Ingresos (1.0 unidad/op)",
         [fmt_short_mxn(t.ingresos_por_escenario[1.0]) for t in plan_output.trimestres]),
        ("Margen acumulado",
         [fmt_short_mxn(t.margen_por_escenario[1.0]) for t in plan_output.trimestres]),
    ]

    n_rows = 1 + len(filas)
    table = doc.add_table(rows=n_rows, cols=5)
    table.autofit = False
    col_widths = [Inches(2.2), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2)]
    for i, w in enumerate(col_widths):
        table.columns[i].width = w
    set_table_cant_split(table)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, ZEBRA_BLACK)
        set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_cell_vertical_alignment(cell, 'center')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=8.5, bold=True, color=ZEBRA_YELLOW)

    # Filas
    for r_idx, (concepto, valores) in enumerate(filas, start=1):
        # Concepto (columna 1)
        cell = table.rows[r_idx].cells[0]
        cell.width = col_widths[0]
        set_cell_background(cell, ZEBRA_WHITE)
        set_cell_borders(cell, color=ZEBRA_LIGHT_GREY, size=4)
        set_cell_margins(cell, top=80, bottom=80, left=140, right=120)
        set_cell_vertical_alignment(cell, 'center')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(concepto)
        style_run(r, size=9, bold=True, color=ZEBRA_BLACK)

        # Valores trimestrales (columnas 2-5)
        for c_idx, valor in enumerate(valores, start=1):
            cell_v = table.rows[r_idx].cells[c_idx]
            cell_v.width = col_widths[c_idx]
            set_cell_background(cell_v, ZEBRA_WHITE)
            set_cell_borders(cell_v, color=ZEBRA_LIGHT_GREY, size=4)
            set_cell_margins(cell_v, top=80, bottom=80, left=120, right=120)
            set_cell_vertical_alignment(cell_v, 'center')

            cell_v.paragraphs[0].clear()
            p = cell_v.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            r = p.add_run(valor)
            style_run(r, size=9, color=ZEBRA_DARK_GREY)

    add_spacer(doc, 8)

    # Caja de totales 12 meses (fondo negro estilo KPI)
    add_styled_paragraph(
        doc, "RESUMEN ANUAL: TOTALES 12 MESES",
        size=9, bold=True, color=ZEBRA_MID_GREY,
        space_before=2, space_after=4
    )

    totales = [
        ("INVERSIÓN TOTAL", fmt_short_mxn(plan_output.inversion_total_anual)),
        ("CIERRES REALES", fmt_int(plan_output.cierres_reales_anual)),
        ("INGRESOS PROYECTADOS", fmt_short_mxn(plan_output.ingresos_anual_por_escenario[1.0])),
        ("ROA ANUAL", fmt_roa(plan_output.roa_anual_por_escenario[1.0])),
    ]

    tot_table = doc.add_table(rows=1, cols=4)
    tot_table.autofit = False
    tcol_w = Inches(7.0 / 4)
    for col in tot_table.columns:
        col.width = tcol_w
    set_table_cant_split(tot_table)

    for i, (label, valor) in enumerate(totales):
        cell = tot_table.rows[0].cells[i]
        cell.width = tcol_w
        set_cell_background(cell, ZEBRA_BLACK)
        set_cell_borders(cell, color=ZEBRA_BLACK, size=4)
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        set_cell_vertical_alignment(cell, 'center')

        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        r = p.add_run(label)
        style_run(r, size=7.5, bold=True, color=ZEBRA_MID_GREY)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)
        r2 = p2.add_run(valor)
        style_run(r2, size=15, bold=True, color=ZEBRA_YELLOW)

    add_spacer(doc, 6)

    # Nota al pie
    nota = (
        "Proyección base con 1 unidad por operación. El Excel anexo incluye escenarios "
        "1.2x / 1.4x / 1.6x para casos donde el cliente compra más de una unidad por "
        "operación, así como detalle mensual completo y todos los inputs editables."
    )
    add_styled_paragraph(
        doc, nota,
        size=8.5, italic=True, color=ZEBRA_MID_GREY,
        space_before=0, space_after=8, line_spacing=1.4
    )


def build_signature(doc, firma="Zebra High Performance Marketing"):
    """Firma final del documento."""
    add_styled_paragraph(
        doc, firma,
        size=9, bold=True, color=ZEBRA_MID_GREY,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=12, space_after=0
    )


# =============================================================================
# FUNCIÓN PRINCIPAL
# =============================================================================

def generate_proposal(data, output_path):
    """
    Genera una propuesta estratégica de Zebra como archivo DOCX.

    Esquema de `data`:
    {
        "titulo_propuesta": "PROPUESTA ESTRATÉGICA",
        "subtitulo_propuesta": "Nombre Cliente",

        "cliente": "Persona — Empresa",
        "modelo": "Modelo dominante + soportes",
        "objetivo": "Cifra · ubicación",

        "contexto": {
            "intro": "Párrafo introductorio del negocio.",
            "proyectos_titulo": "LOS PROYECTOS",
            "proyectos_items": ["Item 1", "Item 2", ...],
            "situacion_titulo": "LA SITUACIÓN HOY",
            "situacion_texto": "Descripción de la situación actual."
        },

        "diagnostico": [
            {"titulo": "Verdad incómoda 1", "texto": "Explicación."},
            ...  # entre 3 y 4 diagnósticos
        ],

        "cambio_paradigma": {
            "antes_titulo": "Lo que se tiene hoy:",
            "antes_texto": "Descripción concreta de la situación.",
            "despues_titulo": "Lo que [cliente] necesita:",
            "despues_texto": "Descripción del sistema necesario.",
            "punchline": "Frase contundente que encapsula el cambio."
        },

        "sistema_propuesto": {
            "intro": "Este caso requiere [N] modelos trabajando juntos:",
            "modelos": [
                {"nombre": "MODELO", "descripcion": "Función específica."},
                ...  # 2 a 4 modelos
            ]
        },

        "arquitectura": {
            "modo": "capas" o "flat",
            # Si modo="capas":
            "capas": [
                {
                    "nombre": "CAPA 1 — NOMBRE",
                    "items": [
                        {"numero": 1, "titulo": "...", "bullets": ["...", "..."]},
                        ...
                    ]
                },
                ...
            ],
            # Si modo="flat":
            "items": [
                {"numero": 1, "titulo": "...", "bullets": [...]},
                ...
            ]
        },

        "contenido": {  # Opcional — omitir si el sistema no requiere contenido
            "intro": "Introducción de la sección.",
            "piezas": [
                {"numero": 1, "pieza": "Nombre de pieza", "funcion": "Función."},
                ...
            ]
        },

        "plan_implementacion": [
            {
                "fase": "FASE 1",
                "semanas": "Semanas 1–3",
                "titulo": "Nombre de la fase",
                "actividades": ["Actividad 1", ...]
            },
            ...  # típicamente 4 fases
        ],

        "metas": {
            "kpis": [
                {"valor": "> 70%", "descripcion": "Descripción corta"},
                ...  # 3 KPIs
            ],
            "escenario_base": {
                "titulo": "ESCENARIO BASE — PRIMER TRIMESTRE",
                "subtitulo": "Con X + Y + Z:",
                "bullets": ["...", "..."],
                "cierre": "Frase honesta sobre timing."
            }
        },

        "inversion": {
            "setup_titulo": "SETUP INICIAL",
            "setup_monto": "$60,000 MXN",
            "setup_items": ["...", ...],
            "fee_titulo": "FEE MENSUAL",
            "fee_monto": "$90,000 MXN",
            "fee_items": ["...", ...],
            "pauta_texto": "Descripción de la lógica de pauta.",
            "pauta_fase_inicial": "$30,000–$50,000 MXN / mes ..."
        },

        "por_que_zebra": {
            "agencia_normal": "Descripción.",
            "zebra": "Descripción.",
            "cierre": "No somos la agencia... Somos el sistema..."
        },

        "firma": "Zebra High Performance Marketing"  # Opcional
    }
    """
    doc = Document()
    set_page_setup(doc)
    set_default_document_font(doc)

    # Título
    build_title_block(
        doc,
        titulo=data.get('titulo_propuesta', 'PROPUESTA ESTRATÉGICA'),
        subtitulo=data.get('subtitulo_propuesta', '')
    )

    # Header tripartita
    build_header_tripartita(
        doc,
        cliente=data['cliente'],
        modelo=data['modelo'],
        objetivo=data['objetivo']
    )

    # 01 — Contexto
    build_section_header(doc, "01", "CONTEXTO DEL CLIENTE")
    ctx = data['contexto']
    build_context_section(
        doc,
        intro=ctx['intro'],
        proyectos_titulo=ctx['proyectos_titulo'],
        proyectos_items=ctx['proyectos_items'],
        situacion_titulo=ctx['situacion_titulo'],
        situacion_texto=ctx['situacion_texto']
    )

    # 02 — Diagnóstico
    build_section_header(doc, "02", "DIAGNÓSTICO")
    build_diagnosis_boxes(doc, data['diagnostico'])

    # 03 — Cambio de paradigma
    build_section_header(doc, "03", "CAMBIO DE PARADIGMA")
    cp = data['cambio_paradigma']
    build_paradigm_shift(
        doc,
        antes_titulo=cp['antes_titulo'],
        antes_texto=cp['antes_texto'],
        despues_titulo=cp['despues_titulo'],
        despues_texto=cp['despues_texto'],
        punchline=cp['punchline']
    )

    # 04 — Sistema propuesto
    build_section_header(doc, "04", "EL SISTEMA PROPUESTO")
    sp = data['sistema_propuesto']
    build_system_models(doc, intro=sp['intro'], modelos=sp['modelos'])

    # 05 — Arquitectura
    build_section_header(doc, "05", "ARQUITECTURA DEL SISTEMA")
    arq = data['arquitectura']
    if arq.get('modo', 'flat') == 'capas':
        build_architecture(doc, arq['capas'])
    else:
        build_architecture_flat(doc, arq['items'])

    # 06 — Piezas de contenido (opcional)
    section_num = 6
    if 'contenido' in data and data['contenido']:
        build_section_header(doc, f"{section_num:02d}", "IDEAS DE CONTENIDOS")
        build_content_pieces(
            doc,
            intro=data['contenido']['intro'],
            piezas=data['contenido']['piezas']
        )
        section_num += 1

    # 07 — Plan de implementación
    build_section_header(doc, f"{section_num:02d}", "PLAN DE IMPLEMENTACIÓN")
    build_implementation_phases(doc, data['plan_implementacion'])
    section_num += 1

    # 08 — Objetivo / Metas
    build_section_header(doc, f"{section_num:02d}", "OBJETIVO DEL SISTEMA / METAS")
    metas = data['metas']
    build_kpi_row(doc, metas['kpis'])
    eb = metas['escenario_base']
    build_base_scenario(
        doc,
        titulo=eb['titulo'],
        subtitulo=eb['subtitulo'],
        bullets=eb['bullets'],
        cierre=eb['cierre']
    )
    section_num += 1

    # 09 — Inversión
    build_section_header(doc, f"{section_num:02d}", "INVERSIÓN")
    inv = data['inversion']
    build_investment_section(
        doc,
        setup_titulo=inv['setup_titulo'],
        setup_monto=inv['setup_monto'],
        setup_items=inv['setup_items'],
        fee_titulo=inv['fee_titulo'],
        fee_monto=inv['fee_monto'],
        fee_items=inv['fee_items'],
        pauta_texto=inv['pauta_texto'],
        pauta_fase_inicial=inv.get('pauta_fase_inicial', '')
    )

    # 09 (continuación) — Calculadora de Inversión (opcional, solo si hay datos completos)
    if 'calculadora_inversion' in data and data['calculadora_inversion']:
        build_investment_calculator_box(
            doc,
            calc_output=data['calculadora_inversion'],
            anexo_excel_nota=data.get('genera_excel_anexo', True)
        )
    section_num += 1

    # 10 — Proyección a 12 meses (opcional, solo si hay datos completos)
    if ('plan_12_meses' in data and data['plan_12_meses']
            and 'calculadora_inversion' in data and data['calculadora_inversion']):
        build_section_header(doc, f"{section_num:02d}", "PROYECCIÓN A 12 MESES")
        build_12_month_projection(
            doc,
            plan_output=data['plan_12_meses'],
            calc_output=data['calculadora_inversion']
        )
        section_num += 1

    # 11 — Por qué Zebra
    build_section_header(doc, f"{section_num:02d}", "POR QUÉ ZEBRA")
    pz = data['por_que_zebra']
    build_why_zebra(
        doc,
        agencia_normal=pz['agencia_normal'],
        zebra_texto=pz['zebra'],
        cierre=pz['cierre']
    )
    section_num += 1

    # 12 — Cierre (opcional)
    if 'cierre_final' in data and data['cierre_final']:
        build_section_header(doc, f"{section_num:02d}", "CIERRE")
        build_final_closing(doc, data['cierre_final'])

    # Firma
    build_signature(doc, data.get('firma', 'Zebra High Performance Marketing'))

    # Guardar
    doc.save(output_path)
    return output_path
